#!/usr/bin/env python3
# =====================================================================
# ECAS + FrSBe cognitive-behavioural phenotyping in ALS
# BUILD MANUSCRIPT TABLES 2, 3, 4 AS A WORD DOCUMENT
#
# Reads the results/*.csv files produced by:
#   reproduce_analyses.py            -> c9orf72_odds_ratios.csv,
#                                        c9orf72_prevalence_by_class.csv
#   reproduce_behavioural_tables.py  -> table2_table3_values.csv
#
# and writes a single Word file (tables_2_3_4.docx) containing properly
# formatted Table 2, Table 3, and Table 4, styled to match Table 1 in
# the manuscript (Times New Roman, light-blue header shading, thin grey
# borders).
#
# NOTE: Table 1 (cohort characteristics) and Table 5 (detailed genetics
# incl. SOD1/TARDBP) are NOT produced by this script, because the
# underlying per-patient values (age, ALSFRS-R, education, SOD1/TARDBP
# carrier counts, etc.) are not present in the reduced results/*.csv
# files used here -- only the C9orf72 and behavioural-table numbers are.
#
# p-value formatting (fmt_p) and phenotype naming now come from
# als_common.py, so numbers reported here can never drift out of sync
# with the same values reported by make_tables_figures.py.
#
# ENVIRONMENT: python-docx, pandas
# =====================================================================

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from als_common import RESULTS_DIR, fmt_p, PHENOTYPE_SNAKE
import os

OUT_DOCX = "results/tables/tables_2_3_4.docx"

HEADER_FILL = "D9E2F3"   # same light blue as Table 1 header
BORDER_COLOR = "BBBBBB"  # same thin grey border as Table 1
FONT_NAME = "Times New Roman"
FONT_SIZE = 12


# ---------------------------------------------------------------------
# Low-level helpers to replicate the manuscript's table styling exactly
# ---------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_COLOR, sz=2):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def style_cell_text(cell, text, bold=False, size=FONT_SIZE):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        style_cell_text(hdr_cells[i], h, bold=True)
        set_cell_shading(hdr_cells[i], HEADER_FILL)
        set_cell_borders(hdr_cells[i])

    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            style_cell_text(cells[i], str(val), bold=False)
            set_cell_borders(cells[i])

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.bold = True
    return p


# ---------------------------------------------------------------------
# Table 2 & Table 3 (from table2_table3_values.csv)
# ---------------------------------------------------------------------
def build_table2_table3(doc, df):
    classes = PHENOTYPE_SNAKE  # ["Preserved", "Cognitive_only", "Behavioural_only", "Severe_FTD"]
    display_names = ["Preserved", "Cognitive-only", "Behavioural-only", "Severe/FTD"]

    for table_name in df["table"].unique():
        sub = df[df["table"] == table_name]
        add_caption(doc, f"Table — {table_name} (median [IQR])")
        headers = ["Measure"] + display_names + ["Kruskal-Wallis p"]
        rows = []
        for _, r in sub.iterrows():
            row = [r["measure"]]
            for c in classes:
                med = r[f"{c}_median"]
                q1 = r[f"{c}_Q1"]
                q3 = r[f"{c}_Q3"]
                row.append(f"{med:.1f} [{q1:.1f}\u2013{q3:.1f}]")
            row.append(fmt_p(r["kruskal_p"]))
            rows.append(row)
        add_styled_table(doc, headers, rows,
                          col_widths=[2.0, 1.3, 1.3, 1.3, 1.3, 1.0])
        doc.add_paragraph()


# ---------------------------------------------------------------------
# Table 4 (C9orf72 odds ratios + BCH-weighted prevalence by class)
# ---------------------------------------------------------------------
def build_table4(doc, or_df, prev_df):
    add_caption(doc, "Table 4 — C9orf72 status by phenotype (BCH-adjusted)")

    headers = ["Phenotype", "C9orf72 prevalence (%)"]
    rows = [[r["class"], f"{r['C9orf72_prevalence_pct']:.1f}"] for _, r in prev_df.iterrows()]
    add_styled_table(doc, headers, rows, col_widths=[2.5, 2.0])
    doc.add_paragraph()

    headers2 = ["Comparison", "OR", "95% CI", "p"]
    rows2 = []
    for _, r in or_df.iterrows():
        rows2.append([
            r["comparison"],
            f"{r['OR']:.1f}",
            f"{r['CI_low']:.1f}\u2013{r['CI_high']:.1f}",
            fmt_p(r["p"]),
        ])
    add_styled_table(doc, headers2, rows2, col_widths=[2.8, 1.0, 1.5, 1.0])
    doc.add_paragraph()


# ---------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------
def main():
    tbl23_path = os.path.join(RESULTS_DIR, "table2_table3_values.csv")
    or_path = os.path.join(RESULTS_DIR, "c9orf72_odds_ratios.csv")
    prev_path = os.path.join(RESULTS_DIR, "c9orf72_prevalence_by_class.csv")

    missing = [p for p in (tbl23_path, or_path, prev_path) if not os.path.exists(p)]
    if missing:
        raise FileNotFoundError(
            f"Missing input file(s): {missing}. Run reproduce_analyses.py and "
            f"reproduce_behavioural_tables.py first (they populate results/)."
        )

    df23 = pd.read_csv(tbl23_path)
    or_df = pd.read_csv(or_path)
    prev_df = pd.read_csv(prev_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)

    doc.add_heading("Manuscript tables generated from results/*.csv", level=1)
    doc.add_paragraph(
        "Table 1 (cohort characteristics) and Table 5 (detailed genetics, "
        "SOD1/TARDBP) are NOT included here: the underlying per-patient "
        "values are not present in the reduced results/*.csv files used "
        "by this script."
    )

    build_table2_table3(doc, df23)
    build_table4(doc, or_df, prev_df)

    doc.save(OUT_DOCX)
    print(f"[saved] {OUT_DOCX}")


if __name__ == "__main__":
    main()